"""社内規程集をAI検索インデックスに取り込む(AI Q&Aの人事・総務アシスタント化)。

対象: KITEI_DIR (既定: Y:\\管理本部\\人事課\\高千穂システム規程集\\改定260711(一般))
      配下の .docx / .pdf / .txt
      ※「(一般)」= 全社員公開の規程のみを置くフォルダ。機密規程は入れないこと。
方式: テキスト抽出 → 条単位を目安にチャンク分割 → photo-index に
      media_type="kitei" で全量入れ替えupsert(古い規程チャンクは先に削除)。

実行(Yにアクセスできるデスクトップで、規程改定時に1回):
    python export_kitei_index.py            # 取り込み
    python export_kitei_index.py --dry-run  # 対象と分割数の確認のみ
    (または run_kitei_index.bat をダブルクリック)

依存: pip install python-docx pypdf (requirements.txt に追加済み)
"""
from __future__ import annotations

import argparse
import hashlib
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).with_name(".env"), encoding="utf-8")
except Exception:
    pass

sys.path.insert(0, str(Path(__file__).resolve().parent))

KITEI_DIR = Path(os.environ.get(
    "KITEI_DIR",
    r"Y:\管理本部\人事課\高千穂システム規程集\改定260711(一般)",
))
MEDIA_TYPE = "kitei"
CHUNK_MAX = 1200      # 1チャンクの目安文字数
BATCH = 100

_JOU = re.compile(r"(?=第\s*[0-9０-９一二三四五六七八九十百]+\s*条)")


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            print("[ERROR] python-docx がありません。pip install python-docx", file=sys.stderr)
            sys.exit(1)
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            print("[ERROR] pypdf がありません。pip install pypdf", file=sys.stderr)
            sys.exit(1)
        try:
            return "\n".join((pg.extract_text() or "") for pg in PdfReader(str(path)).pages)
        except Exception as e:
            print(f"[WARN] PDF抽出失敗: {path.name} ({e})", file=sys.stderr)
            return ""
    if ext == ".txt":
        for enc in ("utf-8-sig", "cp932", "utf-8"):
            try:
                return path.read_text(encoding=enc)
            except Exception:
                continue
    return ""


def _chunk(text: str) -> list[str]:
    """「第○条」を目安に分割し、長すぎる塊はさらに分ける。"""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    parts = [p.strip() for p in _JOU.split(text) if p.strip()]
    chunks: list[str] = []
    buf = ""
    for p in parts:
        if len(buf) + len(p) <= CHUNK_MAX:
            buf = f"{buf}\n{p}".strip()
        else:
            if buf:
                chunks.append(buf)
            while len(p) > CHUNK_MAX:  # 条自体が長い場合は固定長で切る
                chunks.append(p[:CHUNK_MAX])
                p = p[CHUNK_MAX - 100:]  # 100文字重ねて文脈を保つ
            buf = p
    if buf:
        chunks.append(buf)
    return chunks


def main() -> None:
    ap = argparse.ArgumentParser(description="規程集をAI検索インデックスへ取り込む")
    ap.add_argument("--dry-run", action="store_true", help="アップロードせず対象を表示")
    args = ap.parse_args()

    if not KITEI_DIR.is_dir():
        print(f"[ERROR] 規程フォルダが見つかりません: {KITEI_DIR}", file=sys.stderr)
        print("        Y: 未接続の可能性。Yにアクセスできる端末で実行してください。", file=sys.stderr)
        sys.exit(1)

    files = sorted(
        p for p in KITEI_DIR.rglob("*")
        if p.is_file() and p.suffix.lower() in {".docx", ".pdf", ".txt"}
        and not p.name.startswith("~$")
    )
    if not files:
        print(f"[ERROR] 対象ファイルがありません: {KITEI_DIR}", file=sys.stderr)
        sys.exit(1)
    print(f"[SCAN] 対象 {len(files)} ファイル: {KITEI_DIR}")

    indexed_at = datetime.now(tz=timezone.utc).isoformat()
    docs: list[dict] = []
    for f in files:
        title = f.stem  # 規程名(ファイル名から)
        text = _extract_text(f)
        chunks = _chunk(text)
        print(f"  - {f.name}: {len(text)}文字 → {len(chunks)}チャンク")
        for i, ch in enumerate(chunks):
            docs.append({
                "id": hashlib.sha256(f"kitei::{f}::{i}".encode("utf-8")).hexdigest(),
                "file_path": str(f),
                "file_name": f.name,
                "workno": "",
                "workno_name": title,          # AI Q&Aの参照表示に使う規程名
                "phase": "",
                "media_type": MEDIA_TYPE,
                "capture_date": None,
                "capture_date_raw": "",
                "extension": f.suffix.lower(),
                "folder_path": str(f.parent),
                "indexed_at": indexed_at,
                "content_text": f"【{title}】\n{ch}",
            })
    print(f"[CHUNK] 合計 {len(docs)} チャンク")
    if args.dry_run:
        print("[DRY-RUN] アップロードせず終了")
        return

    from azure.core.credentials import AzureKeyCredential
    from azure.search.documents import SearchClient
    from rag.config import SEARCH_INDEX_NAME, ensure_search_credentials

    endpoint, api_key = ensure_search_credentials()
    client = SearchClient(endpoint, SEARCH_INDEX_NAME, AzureKeyCredential(api_key))

    # 既存の規程チャンクを全削除(改定での差し替え・削除に対応)
    old_ids = [r["id"] for r in client.search(
        search_text="*", filter=f"media_type eq '{MEDIA_TYPE}'",
        select=["id"], top=100000)]
    if old_ids:
        for i in range(0, len(old_ids), BATCH):
            client.delete_documents([{"id": x} for x in old_ids[i:i + BATCH]])
        print(f"[DELETE] 旧規程チャンク {len(old_ids)} 件を削除")

    for i in range(0, len(docs), BATCH):
        client.merge_or_upload_documents(docs[i:i + BATCH])
    print(f"[UPLOAD] {len(docs)} チャンクを登録完了")
    print("[DONE] AI Q&Aで規程を参照できるようになりました。")


if __name__ == "__main__":
    main()
