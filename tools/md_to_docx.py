import os
import re
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_style_font(style, font_name, font_size=None, bold=None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_size is not None:
        style.font.size = Pt(font_size)
    if bold is not None:
        style.font.bold = bold


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Meiryo", 10.5)

    for level, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = doc.styles[level]
        set_style_font(style, "Meiryo", size, True)


def add_paragraph_with_indent(doc, text, left_cm=0.0, bullet=None, bold=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(left_cm)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"{bullet} {text}" if bullet else text)
    run.font.name = "Meiryo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")
    run.font.size = Pt(10.5)
    run.bold = bold
    return paragraph


def convert_md_to_docx(md_text, out_path):
    doc = Document()
    configure_document(doc)

    in_code = False
    for raw in md_text.splitlines():
        line = raw.rstrip("\n")

        if line.startswith("```"):
            in_code = not in_code
            if not in_code:
                doc.add_paragraph()
            continue

        if in_code:
            add_paragraph_with_indent(doc, line, left_cm=1.0)
            continue

        if not line.strip():
            doc.add_paragraph()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1:
                paragraph = doc.add_paragraph(style="Title")
                run = paragraph.add_run(text)
                run.font.underline = True
                run.font.name = "Meiryo"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")
            else:
                heading_level = min(level - 1, 3)
                paragraph = doc.add_paragraph(style=f"Heading {heading_level}")
                run = paragraph.add_run(text)
                run.font.name = "Meiryo"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")
                if level <= 4:
                    run.font.underline = True
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(4)
            continue

        stripped = line.lstrip()
        leading_spaces = len(line) - len(stripped)
        indent_cm = 0.0
        if leading_spaces >= 4:
            indent_cm = 1.6
        elif leading_spaces >= 2:
            indent_cm = 0.8

        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered_match:
            add_paragraph_with_indent(doc, ordered_match.group(2), left_cm=indent_cm, bullet=f"{ordered_match.group(1)}.")
            continue

        if stripped.startswith("- "):
            add_paragraph_with_indent(doc, stripped[2:], left_cm=indent_cm, bullet="・")
            continue

        add_paragraph_with_indent(doc, stripped, left_cm=indent_cm)

    doc.save(out_path)


def main(md_path, out_path):
    if not os.path.exists(md_path):
        print("MD not found:", md_path)
        return 2

    with open(md_path, "r", encoding="utf-8") as handle:
        md_text = handle.read()

    convert_md_to_docx(md_text, out_path)
    print("Wrote DOCX:", out_path)
    return 0


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py input.md output.docx")
        sys.exit(2)
    sys.exit(main(sys.argv[1], sys.argv[2]))